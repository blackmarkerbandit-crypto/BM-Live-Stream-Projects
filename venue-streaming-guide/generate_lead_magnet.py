"""
Generate a professional lead-magnet Word document for venue streaming services.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# -- Page setup --
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# -- Color constants --
BRAND_DARK = RGBColor(0x1A, 0x1A, 0x2E)   # Dark navy
BRAND_PRIMARY = RGBColor(0x00, 0x7A, 0xCC)  # Blue
BRAND_ACCENT = RGBColor(0xE8, 0x5D, 0x26)   # Orange
BRAND_LIGHT = RGBColor(0xF5, 0xF5, 0xF5)    # Light gray
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x33, 0x33, 0x33)
GRAY = RGBColor(0x66, 0x66, 0x66)

# -- Style helpers --
def set_cell_shading(cell, color_hex):
    """Set background color on a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_paragraph(doc_or_cell, text, font_size=11, bold=False, color=BLACK,
                         alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6, space_before=0,
                         italic=False, font_name="Calibri"):
    p = doc_or_cell.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.italic = italic
    run.font.name = font_name
    return p

def add_bullet(doc, text, font_size=11, color=BLACK, bold_prefix="", indent_level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(1)
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.font.size = Pt(font_size)
        run_b.font.bold = True
        run_b.font.color.rgb = color
        run_b.font.name = "Calibri"
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.color.rgb = color
    run.font.name = "Calibri"
    return p

def add_section_heading(doc, text):
    add_styled_paragraph(doc, "", font_size=6, space_after=0)
    p = add_styled_paragraph(doc, text, font_size=18, bold=True, color=BRAND_PRIMARY,
                             space_before=12, space_after=6)
    # Add a thin line under the heading
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="4" w:color="007ACC"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p

def make_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        set_cell_shading(cell, "1A1A2E")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = WHITE
        run.font.name = "Calibri"

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F0F4F8")
            p = cell.paragraphs[0]
            run = p.add_run(str(value))
            run.font.size = Pt(10)
            run.font.color.rgb = BLACK
            run.font.name = "Calibri"

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    return table


# ============================================================
# COVER PAGE
# ============================================================

# Spacer
for _ in range(4):
    add_styled_paragraph(doc, "", font_size=12, space_after=0)

# Title
add_styled_paragraph(
    doc,
    "THE VENUE STREAMING",
    font_size=36, bold=True, color=BRAND_DARK,
    alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0
)
add_styled_paragraph(
    doc,
    "PLAYBOOK",
    font_size=42, bold=True, color=BRAND_PRIMARY,
    alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12
)

# Subtitle
add_styled_paragraph(
    doc,
    "How Smart Venues Are Turning Live Events\nInto Global Experiences",
    font_size=16, color=GRAY,
    alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24, italic=True
)

# Divider line
p_div = doc.add_paragraph()
p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_div = p_div.add_run("_" * 50)
run_div.font.color.rgb = BRAND_PRIMARY
run_div.font.size = Pt(10)

# Value props
for _ in range(2):
    add_styled_paragraph(doc, "", font_size=8, space_after=0)

props = [
    "Reach audiences beyond your four walls",
    "Create new revenue streams from every event",
    "Build a global brand without a global budget",
]
for prop in props:
    add_styled_paragraph(
        doc, f"\u2713  {prop}",
        font_size=14, color=BRAND_DARK,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8
    )

for _ in range(4):
    add_styled_paragraph(doc, "", font_size=10, space_after=0)

# CTA teaser
add_styled_paragraph(
    doc,
    "A free guide from [Your Company Name]",
    font_size=12, color=GRAY, italic=True,
    alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4
)
add_styled_paragraph(
    doc,
    "[your-website.com]  |  [your@email.com]  |  [phone number]",
    font_size=11, color=BRAND_PRIMARY,
    alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0
)

# Page break
doc.add_page_break()


# ============================================================
# TABLE OF CONTENTS
# ============================================================

add_styled_paragraph(
    doc, "WHAT'S INSIDE", font_size=24, bold=True, color=BRAND_DARK,
    alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18, space_before=24
)

toc_items = [
    ("01", "The Live Streaming Opportunity", "Why venues that don't stream are leaving money on the table"),
    ("02", "5 Revenue Streams You're Missing", "Monetization models that work right now"),
    ("03", "What It Actually Takes", "Equipment, internet, and crew — simplified"),
    ("04", "The 3-Tier Approach", "Start small, prove ROI, then scale"),
    ("05", "Platform Strategy", "Where to stream and why it matters"),
    ("06", "Audio: The Make-or-Break Factor", "Why 90% of stream quality is sound"),
    ("07", "Legal Essentials", "Music rights, licensing, and protecting your venue"),
    ("08", "Your Next Step", "How we make this effortless for you"),
]

for num, title, desc in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run_num = p.add_run(f"  {num}   ")
    run_num.font.size = Pt(14)
    run_num.font.bold = True
    run_num.font.color.rgb = BRAND_PRIMARY
    run_num.font.name = "Calibri"
    run_title = p.add_run(f"{title}\n")
    run_title.font.size = Pt(13)
    run_title.font.bold = True
    run_title.font.color.rgb = BRAND_DARK
    run_title.font.name = "Calibri"
    run_desc = p.add_run(f"         {desc}")
    run_desc.font.size = Pt(10)
    run_desc.font.color.rgb = GRAY
    run_desc.font.italic = True
    run_desc.font.name = "Calibri"

doc.add_page_break()


# ============================================================
# SECTION 1: THE OPPORTUNITY
# ============================================================

add_section_heading(doc, "01  The Live Streaming Opportunity")

add_styled_paragraph(
    doc,
    "The live events industry has fundamentally changed. Audiences now expect "
    "a digital option — and the venues delivering it are winning.",
    font_size=12, color=BLACK, space_after=12
)

add_styled_paragraph(doc, "The Numbers Tell the Story", font_size=14, bold=True, color=BRAND_DARK, space_after=8)

make_table(doc,
    ["Metric", "Data Point"],
    [
        ["Global live streaming market", "$100+ billion by 2027 (Grand View Research)"],
        ["Viewers who prefer hybrid events", "67% want the option to attend virtually (Eventbrite)"],
        ["Revenue lift from adding virtual tickets", "15-35% average increase per event"],
        ["Fan willingness to pay for streams", "72% of live music fans have paid for a stream (MIDiA Research)"],
        ["Average virtual ticket price", "$10-25 for concerts, $5-15 for other events"],
    ],
    col_widths=[2.5, 4.0]
)

add_styled_paragraph(doc, "", font_size=6)

add_styled_paragraph(
    doc,
    "Every show you don't stream is an audience you're not reaching and revenue you're not earning.",
    font_size=12, bold=True, color=BRAND_ACCENT, italic=True,
    alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12
)

add_styled_paragraph(doc, "Who's Already Doing This?", font_size=14, bold=True, color=BRAND_DARK, space_after=8)

add_styled_paragraph(
    doc,
    "Forward-thinking venues across every category are streaming today:",
    font_size=11, space_after=6
)

examples = [
    ("Music venues ", "— selling virtual tickets alongside physical ones, reaching fans in other cities and countries"),
    ("Theaters ", "— offering accessible viewing for patrons who can't attend in person"),
    ("Houses of worship ", "— connecting with congregants who are homebound, traveling, or remote"),
    ("Conference centers ", "— hybrid events are now the default expectation for corporate clients"),
    ("Comedy clubs ", "— building national followings from local stages"),
]
for bold_part, rest in examples:
    add_bullet(doc, rest, bold_prefix=bold_part)

doc.add_page_break()


# ============================================================
# SECTION 2: REVENUE STREAMS
# ============================================================

add_section_heading(doc, "02  5 Revenue Streams You're Missing")

add_styled_paragraph(
    doc,
    "Streaming isn't just about exposure — it's a profit center. Here are five proven models:",
    font_size=12, space_after=12
)

revenue_models = [
    (
        "Virtual Ticket Sales (Pay-Per-View)",
        "Sell access to your stream just like a physical ticket. Virtual tickets typically price at "
        "30-50% of in-person tickets. A 500-capacity venue selling 200 virtual tickets at $15 each "
        "adds $3,000 per event with zero additional venue costs.",
        "$3,000-15,000/event"
    ),
    (
        "Subscriptions & Memberships",
        "Charge $10-25/month for access to all your streamed events. With 200 subscribers, that's "
        "$2,000-5,000/month in predictable recurring revenue — even during slow booking periods.",
        "$2,000-5,000/month"
    ),
    (
        "Tips & Donations",
        "Viewers send voluntary payments during streams. Music streams commonly generate $200-1,000+ "
        "per event in tips. Setup takes 10 minutes with tools like Streamlabs or Ko-fi.",
        "$200-1,000/event"
    ),
    (
        "Sponsorships",
        "Local businesses pay to have their brand on your stream — logo placement, branded segments, "
        "pre-roll ads. A single local sponsor can cover your entire streaming operation cost.",
        "$500-5,000/month"
    ),
    (
        "VOD & Content Licensing",
        "Your archived streams become a content library. License performances to artists, sell VOD "
        "replays, or generate YouTube ad revenue. Content works for you long after the show ends.",
        "Ongoing passive income"
    ),
]

for i, (title, desc, revenue) in enumerate(revenue_models):
    add_styled_paragraph(doc, "", font_size=3, space_after=0)

    # Revenue model box
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    left = t.rows[0].cells[0]
    right = t.rows[0].cells[1]
    left.width = Inches(5.0)
    right.width = Inches(1.5)

    set_cell_shading(right, "007ACC")

    left.text = ""
    p_title = left.paragraphs[0]
    run_t = p_title.add_run(f"{i+1}. {title}")
    run_t.font.size = Pt(13)
    run_t.font.bold = True
    run_t.font.color.rgb = BRAND_DARK
    run_t.font.name = "Calibri"

    p_desc = left.add_paragraph()
    run_d = p_desc.add_run(desc)
    run_d.font.size = Pt(10)
    run_d.font.color.rgb = BLACK
    run_d.font.name = "Calibri"

    right.text = ""
    p_rev_label = right.paragraphs[0]
    p_rev_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_l = p_rev_label.add_run("POTENTIAL")
    run_l.font.size = Pt(8)
    run_l.font.bold = True
    run_l.font.color.rgb = WHITE
    run_l.font.name = "Calibri"

    p_rev = right.add_paragraph()
    p_rev.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_r = p_rev.add_run(revenue)
    run_r.font.size = Pt(10)
    run_r.font.bold = True
    run_r.font.color.rgb = WHITE
    run_r.font.name = "Calibri"

add_styled_paragraph(doc, "", font_size=6)

# Callout box
t_callout = doc.add_table(rows=1, cols=1)
t_callout.style = "Table Grid"
t_callout.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = t_callout.rows[0].cells[0]
set_cell_shading(cell, "FFF3E0")
cell.text = ""
p_c = cell.paragraphs[0]
p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_c = p_c.add_run("Combined, these revenue streams can add $50,000-200,000+ annually\nto your venue's bottom line.")
run_c.font.size = Pt(13)
run_c.font.bold = True
run_c.font.color.rgb = BRAND_ACCENT
run_c.font.name = "Calibri"

doc.add_page_break()


# ============================================================
# SECTION 3: WHAT IT TAKES
# ============================================================

add_section_heading(doc, "03  What It Actually Takes")

add_styled_paragraph(
    doc,
    "Venue streaming has gotten dramatically simpler and more affordable. Here's an honest "
    "breakdown of what's involved:",
    font_size=12, space_after=12
)

add_styled_paragraph(doc, "Equipment Essentials", font_size=14, bold=True, color=BRAND_DARK, space_after=8)

make_table(doc,
    ["Component", "What It Does", "Starting Cost"],
    [
        ["PTZ Cameras (2-3)", "Remote-controlled cameras mounted in your venue — no operators needed", "$1,600-6,000"],
        ["Video Switcher", "Cuts between camera angles in real time", "$450-1,300"],
        ["Audio Interface", "Captures clean audio from your mixing board", "$100-250"],
        ["Streaming PC or Encoder", "Compresses and sends video to the internet", "$600-1,500"],
        ["Cabling & Mounts", "HDMI/SDI cables, ceiling mounts, adapters", "$200-500"],
    ],
    col_widths=[1.8, 3.2, 1.2]
)

add_styled_paragraph(doc, "", font_size=6)

add_styled_paragraph(doc, "Internet Requirements", font_size=14, bold=True, color=BRAND_DARK, space_after=8)

add_styled_paragraph(
    doc,
    "You need reliable upload speed — not download. Most venues need far less than they think:",
    font_size=11, space_after=6
)

make_table(doc,
    ["Stream Quality", "Upload Speed Needed", "Typical Cost"],
    [
        ["720p (Good)", "5-10 Mbps upload", "Often included in existing business internet"],
        ["1080p (Professional)", "10-20 Mbps upload", "$50-100/month for business fiber"],
        ["4K (Premium)", "25-50 Mbps upload", "$100-200/month for dedicated fiber"],
    ],
    col_widths=[1.8, 2.2, 2.5]
)

add_styled_paragraph(doc, "", font_size=8)

add_styled_paragraph(doc, "Crew Requirements", font_size=14, bold=True, color=BRAND_DARK, space_after=8)

add_styled_paragraph(
    doc,
    "Modern PTZ camera systems dramatically reduce crew needs:",
    font_size=11, space_after=6
)

make_table(doc,
    ["Setup", "Crew Size", "Skill Level"],
    [
        ["2 PTZ cameras + auto-switching", "0 (fully automated)", "Initial setup only"],
        ["2-3 PTZ cameras + live switching", "1 person", "Moderate (trainable in 1-2 sessions)"],
        ["4-6 cameras + graphics", "2-3 people", "Experienced AV / trained staff"],
        ["Full broadcast production", "3-5 people", "Professional crew"],
    ],
    col_widths=[2.5, 1.5, 2.5]
)

add_styled_paragraph(doc, "", font_size=8)

# Investment summary callout
t_inv = doc.add_table(rows=1, cols=1)
t_inv.style = "Table Grid"
t_inv.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = t_inv.rows[0].cells[0]
set_cell_shading(cell, "E8F4FD")
cell.text = ""
p_inv = cell.paragraphs[0]
p_inv.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_inv = p_inv.add_run(
    "Total starting investment: $3,000-8,000\n"
    "That's often recouped within 2-4 streamed events through virtual ticket sales alone."
)
run_inv.font.size = Pt(12)
run_inv.font.bold = True
run_inv.font.color.rgb = BRAND_PRIMARY
run_inv.font.name = "Calibri"

doc.add_page_break()


# ============================================================
# SECTION 4: 3-TIER APPROACH
# ============================================================

add_section_heading(doc, "04  The 3-Tier Approach: Start Small, Scale Smart")

add_styled_paragraph(
    doc,
    "You don't need to build a broadcast studio on day one. The smartest venues follow a "
    "phased approach that proves ROI at each stage before investing more.",
    font_size=12, space_after=14
)

tiers = [
    (
        "TIER 1: PROVE IT  ($2,500-5,000)",
        "1A1A2E",
        [
            "2 PTZ cameras (wide shot + close-up)",
            "Blackmagic ATEM Mini Pro ($450)",
            "Audio feed from your existing FOH board",
            "Stream to YouTube + Facebook (free platforms)",
            "1 person operates everything",
        ],
        "Goal: Validate audience demand. Measure viewership and engagement for 4-6 events. "
        "This tier pays for itself if even 50 people buy a $10 virtual ticket."
    ),
    (
        "TIER 2: MONETIZE  ($8,000-15,000)",
        "007ACC",
        [
            "3-4 PTZ cameras with preset positions",
            "Dedicated streaming PC with vMix or OBS",
            "Professional audio interface + room microphones",
            "Ticketed streaming on Mandolin, Eventive, or Vimeo OTT",
            "Branded graphics, lower thirds, and overlays",
            "1-2 person crew",
        ],
        "Goal: Generate consistent revenue. Add virtual ticketing, sponsorship packages, and "
        "build a subscriber base. Target: $2,000-5,000/month in streaming revenue."
    ),
    (
        "TIER 3: PROFESSIONALIZE  ($20,000-50,000+)",
        "E85D26",
        [
            "5-6 cameras (PTZ + operated) with full coverage",
            "Broadcast-quality switcher with dedicated control panel",
            "Redundant internet with automatic failover",
            "ISO recording of all camera angles",
            "Dedicated stream audio mix with ambient mics",
            "Instant replay capability",
            "2-4 person professional crew",
        ],
        "Goal: Premium product commanding premium pricing. Virtual tickets at $20-50, corporate "
        "hybrid events at premium rates, national sponsorship deals."
    ),
]

for title, color_hex, items, goal in tiers:
    # Tier header
    t_tier = doc.add_table(rows=1, cols=1)
    t_tier.style = "Table Grid"
    t_tier.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t_tier.rows[0].cells[0]
    set_cell_shading(cell, color_hex)
    cell.text = ""
    p_tier = cell.paragraphs[0]
    p_tier.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_tier = p_tier.add_run(title)
    run_tier.font.size = Pt(14)
    run_tier.font.bold = True
    run_tier.font.color.rgb = WHITE
    run_tier.font.name = "Calibri"

    for item in items:
        add_bullet(doc, item, font_size=11)

    add_styled_paragraph(doc, goal, font_size=11, italic=True, color=GRAY, space_after=14, space_before=4)

doc.add_page_break()


# ============================================================
# SECTION 5: PLATFORM STRATEGY
# ============================================================

add_section_heading(doc, "05  Platform Strategy: Where to Stream")

add_styled_paragraph(
    doc,
    "Different platforms serve different goals. Here's where to focus based on what you want to achieve:",
    font_size=12, space_after=12
)

make_table(doc,
    ["Your Goal", "Best Platform", "Why"],
    [
        ["Maximum free reach", "YouTube Live + Facebook Live", "Largest audiences, free, permanent VOD archive"],
        ["Sell virtual tickets", "Mandolin / Eventive / Vimeo OTT", "Built-in ticketing, access control, payment processing"],
        ["Build a loyal community", "Twitch or YouTube", "Subscription models, chat culture, repeat viewership"],
        ["Corporate / branded events", "Vimeo Livestream", "Ad-free, white-label, password protection"],
        ["Cover all bases", "Restream (multi-platform)", "Stream to 30+ platforms from one feed simultaneously"],
    ],
    col_widths=[1.5, 2.2, 2.8]
)

add_styled_paragraph(doc, "", font_size=8)

add_styled_paragraph(
    doc,
    "Our recommendation: Start with YouTube + Facebook (free, maximum reach), then add a "
    "ticketed platform once you've validated demand.",
    font_size=12, bold=True, color=BRAND_PRIMARY, space_after=12
)

doc.add_page_break()


# ============================================================
# SECTION 6: AUDIO
# ============================================================

add_section_heading(doc, "06  Audio: The Make-or-Break Factor")

add_styled_paragraph(
    doc,
    "Here's a truth most people learn the hard way: viewers will tolerate mediocre video "
    "for hours, but they'll leave within 30 seconds of bad audio.",
    font_size=12, space_after=12
)

add_styled_paragraph(doc, "The Right Way to Capture Audio", font_size=14, bold=True, color=BRAND_DARK, space_after=8)

make_table(doc,
    ["Method", "Quality", "Complexity", "Best For"],
    [
        ["Camera microphone", "Poor", "Zero", "Never use this for a real stream"],
        ["Direct feed from FOH board", "Good", "Low", "Quick and easy starting point"],
        ["Dedicated stream mix (aux send)", "Excellent", "Moderate", "Best balance of quality vs. effort"],
        ["Full audio split + stream mixer", "Broadcast", "High", "Large productions, ultimate control"],
    ],
    col_widths=[2.0, 1.0, 1.0, 2.5]
)

add_styled_paragraph(doc, "", font_size=8)

add_styled_paragraph(doc, "The Secret Ingredient: Room Microphones", font_size=14, bold=True, color=BRAND_DARK, space_after=8)

add_styled_paragraph(
    doc,
    "A direct board feed sounds like a studio recording — clean but lifeless. Adding 1-2 room "
    "microphones captures crowd energy, applause, and ambient atmosphere that makes viewers "
    "feel like they're actually there. This one addition transforms stream quality more than "
    "any camera upgrade.",
    font_size=11, space_after=10
)

add_styled_paragraph(
    doc,
    "We handle all of this for you — from the audio interface to the final mix that makes "
    "your stream sound broadcast-ready.",
    font_size=12, bold=True, color=BRAND_PRIMARY, italic=True,
    alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12
)

doc.add_page_break()


# ============================================================
# SECTION 7: LEGAL
# ============================================================

add_section_heading(doc, "07  Legal Essentials You Need to Know")

add_styled_paragraph(
    doc,
    "Music licensing for streaming is different from your standard venue performance license. "
    "Here's what you need to stay protected:",
    font_size=12, space_after=12
)

add_styled_paragraph(doc, "Quick Reference", font_size=14, bold=True, color=BRAND_DARK, space_after=8)

make_table(doc,
    ["Scenario", "Risk Level", "What You Need"],
    [
        ["Band plays original music", "Low", "Written permission from artist (in booking contract)"],
        ["Band plays cover songs", "Moderate", "PRO streaming licenses (ASCAP, BMI, SESAC, GMR)"],
        ["Recorded music during intermission", "High", "Royalty-free music library (Epidemic Sound, Artlist)"],
        ["DJ set / all recorded music", "Very High", "Direct master licenses (complex) — avoid streaming unless cleared"],
        ["Worship services", "Low", "CCLI Streaming License (~$100-500/year)"],
    ],
    col_widths=[2.0, 1.0, 3.5]
)

add_styled_paragraph(doc, "", font_size=8)

add_styled_paragraph(doc, "Essential Protections", font_size=14, bold=True, color=BRAND_DARK, space_after=8)

protections = [
    ("Streaming clause in every booking contract ", "— get written permission from artists to stream their performance"),
    ("PRO licenses for streaming ", "— your existing venue performance license likely doesn't cover online streaming"),
    ("Venue signage ", '— "This event is being live streamed" posted at entrances'),
    ("Royalty-free music ", "— for all non-live segments (pre-show, intermission, post-show)"),
]
for bold_part, rest in protections:
    add_bullet(doc, rest, bold_prefix=bold_part)

add_styled_paragraph(doc, "", font_size=6)

add_styled_paragraph(
    doc,
    "We help our clients navigate all of this — from contract language to PRO registration "
    "to platform compliance.",
    font_size=12, bold=True, color=BRAND_PRIMARY, italic=True,
    alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8
)

doc.add_page_break()


# ============================================================
# SECTION 8: CTA / NEXT STEPS
# ============================================================

add_section_heading(doc, "08  Your Next Step")

for _ in range(2):
    add_styled_paragraph(doc, "", font_size=8, space_after=0)

add_styled_paragraph(
    doc,
    "You've seen the opportunity. You know the numbers.\nNow imagine this for your venue:",
    font_size=14, color=BRAND_DARK,
    alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=16
)

vision_items = [
    "Every event you host reaches 2x-10x more people",
    "A new revenue stream that grows with every show",
    "Artists choosing your venue because you offer streaming",
    "Corporate clients booking you because hybrid events are built in",
    "A growing content library that markets your venue 24/7",
]
for item in vision_items:
    add_styled_paragraph(
        doc, f"\u2713  {item}",
        font_size=13, color=BRAND_DARK,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6
    )

add_styled_paragraph(doc, "", font_size=12)

# CTA box
t_cta = doc.add_table(rows=1, cols=1)
t_cta.style = "Table Grid"
t_cta.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = t_cta.rows[0].cells[0]
set_cell_shading(cell, "1A1A2E")
cell.text = ""

p1 = cell.paragraphs[0]
p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p1.paragraph_format.space_before = Pt(16)
run1 = p1.add_run("LET'S TALK ABOUT YOUR VENUE")
run1.font.size = Pt(20)
run1.font.bold = True
run1.font.color.rgb = WHITE
run1.font.name = "Calibri"

p2 = cell.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run(
    "We'll assess your space, recommend the right setup for your budget,\n"
    "and show you exactly what your streaming revenue could look like."
)
run2.font.size = Pt(12)
run2.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
run2.font.name = "Calibri"

p3 = cell.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_after = Pt(4)
run3 = p3.add_run("FREE VENUE STREAMING ASSESSMENT")
run3.font.size = Pt(16)
run3.font.bold = True
run3.font.color.rgb = BRAND_ACCENT
run3.font.name = "Calibri"

p4 = cell.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.paragraph_format.space_after = Pt(16)
run4 = p4.add_run(
    "[Your Website URL]\n"
    "[your@email.com]\n"
    "[Phone Number]"
)
run4.font.size = Pt(13)
run4.font.bold = True
run4.font.color.rgb = WHITE
run4.font.name = "Calibri"

add_styled_paragraph(doc, "", font_size=12)

# What happens next
add_styled_paragraph(doc, "What Happens When You Reach Out", font_size=14, bold=True, color=BRAND_DARK,
                     alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

steps = [
    ("1. Free consultation call ", "— we learn about your venue, events, and goals (30 min)"),
    ("2. Custom streaming plan ", "— tailored equipment, platform, and monetization recommendations"),
    ("3. Revenue projection ", "— realistic numbers based on your event schedule and audience"),
    ("4. Pilot event ", "— we handle everything for your first stream so you can see it in action"),
]
for bold_part, rest in steps:
    add_bullet(doc, rest, bold_prefix=bold_part)

add_styled_paragraph(doc, "", font_size=10)

add_styled_paragraph(
    doc,
    "No pressure. No obligation. Just a conversation about what's possible.",
    font_size=13, color=GRAY, italic=True,
    alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4
)


# ============================================================
# SAVE
# ============================================================

output_path = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "Venue-Streaming-Playbook.docx"
)
doc.save(output_path)
print(f"Document saved to: {output_path}")
